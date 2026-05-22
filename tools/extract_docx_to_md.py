from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]

SOURCES = [
    (
        Path("/Users/wqq/Downloads/单人拍摄 内容.docx"),
        ROOT / "整理版" / "单人拍摄内容说明.md",
    ),
    (
        Path("/Users/wqq/Downloads/单人拍摄的补拍画面.docx"),
        ROOT / "整理版" / "单人拍摄补拍画面清单.md",
    ),
]


def table_to_md(table):
    rows = []
    for row in table.rows:
        rows.append([cell.text.replace("\n", " / ").strip() for cell in row.cells])
    if not rows:
        return []
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    lines = []
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join(["---"] * width) + " |")
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return lines


def extract_docx(src, dst):
    doc = Document(src)
    lines = [f"# {src.stem}", ""]

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name.startswith("Heading"):
            level = "".join(ch for ch in style_name if ch.isdigit()) or "2"
            level_num = min(max(int(level), 2), 4)
            lines.append("#" * level_num + " " + text)
        else:
            lines.append(text)
        lines.append("")

    if doc.tables:
        lines.append("## 表格内容")
        lines.append("")
        for idx, table in enumerate(doc.tables, start=1):
            lines.append(f"### 表格 {idx}")
            lines.append("")
            lines.extend(table_to_md(table))
            lines.append("")

    dst.parent.mkdir(parents=True, exist_ok=True)
    dst.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def main():
    for src, dst in SOURCES:
        extract_docx(src, dst)


if __name__ == "__main__":
    main()
